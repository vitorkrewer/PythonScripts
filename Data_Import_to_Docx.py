# Nome do Script: DataImport_to_Docx.py
# Autor: Vitor Krewer
# Data de Criação: 14/07/2023
# Descrição: Buscar informações em uma planilha e substituir determinadas marcações em documentos Microsoft Word (.docx). A principal função é automatizar a geração de documentos a partir de informações originárias de uma planlha

'''
Módulos utilizados:
- Python-docx: é uma biblioteca Python para criar e atualizar arquivos do Microsoft Word (.docx).
- Pandas: o pandas do Python é uma biblioteca de código aberto que fornece estruturas de dados e ferramentas de análise de dados de alta performance.
pip install pandas ou pip3 install pandas
pip install python-docx pip3 install python-docx
'''

# -*- coding: utf-8 -*-

from docx import Document
import pandas as pd

# Ler a planilha do Excel com as Informações
df = pd.read_excel('dados_base.xlsx')

# Extrair os valores da coluna 'Informacoes_Antigas = Tag' e 'Novos Trechos = Informação a ser inserida' da planilha. 

'''
O principal objetivo é automatizar a inserção de informação em um documento word com "tags". As tags marcam as informações a serem substituídas. Ao utilizar o script, para facilitar, inclua o nome das colunas utilizando regras de nomenclatura de banco de dados (minha sugestão).

'''

info_tags = df['Informacoes_Antigas'].tolist()
new_infos = df['Novas_Informacoes'].tolist()

'''
df: É uma variável que contém o DataFrame. O DataFrame foi criado lendo um arquivo Excel com a função pd.read_excel('dados_base.xlsx').
tolist(): É um método do Pandas que é aplicado à série (coluna) selecionada, no caso, as especificadas pela variável que contém um DataFrame. Esse método converte a série em uma lista Python. Após aplicar tolist(), a variável info_tags (ou o nome que será utilizado nesta váriavel) conterá uma lista que contém os valores da coluna "Informacoes_Antigas" do DataFrame df. O mesmo ocorrerá com a Novas_Informacoes.
'''

# Abrir o arquivo em formato .docx indicado.
doc = Document("modelo.docx")

# Percorrer as tabelas do documento
for table in doc.tables:
    # Percorrer as linhas da tabela
    for row in table.rows:
        # Percorrer as células da linha
        for cell in row.cells:
            # Verificar cada informação antiga (a ser substituída)
            for info_tags, new_infos in zip(info_tags, new_infos):
                # Verificar se a informação antiga está presente na célula(cell)
                if info_tags in cell.text:
                    # Substituir a informação antiga pela nova informação na célula(cell)
                    modified_text = cell.text.replace(info_tags, new_infos)
                    break

# Salvar as alterações no mesmo arquivo. Caso seja necessário criar um novo arquivo, adicione prefixo ou sulfixo ao nome do arquivo. Exemplo modelo_versao_01.docx

doc.save("modelo_versao1.docx")


'''
Sendo necessário adicionar a leitura de mais linhas e colunas, ou seja, adicionar variáveis 'df que contenham DataFrame', é possível, e necessário: adicionar novas linhas logo abaixo da variável que irá ler as informações da Planilha. Inserir:
outra_info_tags = df['Outra Informação'].tolist()
nova_outra_infos = df['Novas_Informacoes'].tolist()
for info_tags, new_infos, outra_info_tags, nova_outra_infos in zip(info_tags, new_infos, outra_info_tagnova_outra_infos):
    # Verificar se o trecho antigo está presente na célula
    if info_tags in cell.text:
        # Substituir o trecho antigo pelo novo trecho na célula
        modified_text = cell.text.replace(info_tags, new_infos)
        cell.text = modified_text
    if outra_info_tags in cell.text:
        # Substituir a outra informação pelo novo valor na célula
        modified_text = cell.text.replace(outra_info_tags, nova_outra_infos)
        break

No documento .docx utilizei a mesma regra adotada pelo AutoCrat, extensão disponível para o Google Planilhas. Como <<Info1>>, <<Info2>>.

É importante ressaltar que este script tem como função substituir informações de dentro de tabela de um documento do Microsoft Word.
Se o objetivo for a substituição de conteúdo localizado em parágrafos de texto do documento .docx, utilizamos paragraph e paragraph.text.

# Percorrer os parágrafos do documento
for paragraph in doc.paragraphs:
    # Verificar cada trecho antigo
    for new_infos in info_tags:
        # Verificar se o parágrafo contém o trecho antigo a ser substituído
        if new_infos in paragraph.text:
            # Obter o índice do trecho antigo
            index = info_tags.index(new_infos)

            # Obter o novo trecho correspondente da coluna 'Novas Informações' da planilha
            new_infos = df.at[index, 'Novo Trecho']

            # Modificar o texto do parágrafo
            modified_text = paragraph.text.replace(new_infos, new_infos)

            # Atribuir o novo texto ao parágrafo
            paragraph.text = modified_text
'''